import requests
import re
from datetime import datetime, timedelta
from typing import List
import yaml
from docx import Document
from docx.shared import Pt
import json
import copy
from src.exeptions.ai_exceptions.ai_cant_handle_request_exception import AICantHandleRequestException
from src.bitrix.bitrix_manager import BitrixManager
from src.bitrix.bitrix_user import BitrixUser
from src.bitrix.bitrix_task import BitrixTask
from src.AI.identified_names_handler_interface import IdentifiedNamesHandlerInterface
from tzlocal import get_localzone
from src.AI.database.user_database_interface import UserDatabaseInterface

class Deadline():
    time: datetime
    approx_discription: str

    def __init__(self, time: datetime = None, approx_discription: str = "-"):
        self.time = time
        self.approx_discription = approx_discription
    
    def __str__(self):
        if self.time is None:
            return self.approx_discription
        else:
            return self.time.strftime("%Y-%m-%d %H:%M:%S")



class Task:
    def __init__(self, title: str, content: str, deadline: Deadline):
        self.title = title
        self.content = content
        self.deadline = deadline
    
    def __str__(self):
        return f"({self.deadline}) {self.content}"
    
    def __dict__(self) -> dict:
        return {
            "content": self.content,
            "deadline": str(self.deadline)
        }


class Assignee:
    def __init__(self, name: str, original_speaker_name: str, tasks: List[Task]):
        self.name = name
        self.original_speaker_name = original_speaker_name
        self.tasks = tasks
    
    def __str__(self):
        output = f"{self.name}:\n"
        for task in self.tasks:
            output += f"\t{task}\n"
        return output
    
    def __dict__(self) -> dict:
        return {
            "name": self.name,
            "tasks": [task.__dict__() for task in self.tasks]
        }



class EloquityAI:
    def __init__(self, api_key: str, bitrix_manager: BitrixManager, users_database: UserDatabaseInterface, model_name: str = 'gpt-4o'):
        self.api_url = "https://gptunnel.ru/v1/chat/completions"
        self.api_key = api_key
        self.name_identification_prefix = self._load_prefix("prefixes/name_identification.txt")
        self.task_assigment_prefix = self._load_prefix("prefixes/task_assigment.txt")
        self.fix_json_format_prefix = self._load_prefix("prefixes/fix_json_format.txt")
        self.preloaded_names_assigment = self._load_prefix("prefixes/preloaded_names_assigment.txt")
        self.bitrix_finder_prompt = self._load_prefix("prefixes/bitrix_finder_prompt.txt")
        self.model_name = model_name
        self.bitrix_manager = bitrix_manager
        self.users_database = users_database
        self.bitrix_users = self.bitrix_manager.find_users(count_return_entries=-1)
        self.users_database.add_users(self.bitrix_users)
    
    def get_bitrix_users_dict(self, bitrix_users: List[BitrixUser]):
        full_names_to_bitrix_users = {}
        for user in self.bitrix_users:
            full_name = f"{user.last_name} {user.name} {user.second_name}"
            full_names_to_bitrix_users[full_name] = user
        return full_names_to_bitrix_users

    def get_employees_prompt(self):
        employees_prompt = "\n\n".join([str(user) for user in self.bitrix_users])
        return employees_prompt

    def _load_prefix(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except FileNotFoundError:
            print(f"Error: File '{file_path}' not found.")
            return ""

    def get_model_response(self, message):
        headers = {
            "Authorization": self.api_key,
            "Content-Type": "application/json"
        }
        data = {
            "model": self.model_name,
            "max_tokens": 2000,
            "messages": [{"role": "user", "content": message}]
        }

        response = requests.post(self.api_url, headers=headers, json=data)
        response.raise_for_status()  # Raise an error for HTTP errors
        
        content = response.json()["choices"][0]["message"]["content"]

        return content
    
    def get_delta_time_from_str(self, time_string):
        pattern = r"(?:(\d+)d)?(?:(\d+)h)?(?:(\d+)m)?(?:(\d+)s)?"
        
        matches = re.match(pattern, time_string)
        
        if matches:
            days = int(matches.group(1) or 0)
            hours = int(matches.group(2) or 0)
            minutes = int(matches.group(3) or 0)
            seconds = int(matches.group(4) or 0)
            
            return timedelta(days=days, hours=hours, minutes=minutes, seconds=seconds)
        
        return timedelta()
    
    def replace_speakers(self, text, speakers):
        pattern = re.compile(r'\b(speaker_\d+)\b')  # Match words like speaker_0
        return pattern.sub(lambda match: f"[{speakers.get(match.group(1), match.group(1))}]", text)
    
    def paste_date(self, text: str, json_log: dict = None):
        current_date = datetime.now()
        current_time = "Текущая дата: " + current_date.strftime("%H:%M %d.%m.%Y") + "\n"
        texts_modified = text.replace("[CURRENT_DATE]", current_time)

        if json_log is not None:
            json_log["current_time_str"] = current_time

        return texts_modified
    
    def paste_google_meet_nicknames(self, text: str, name_dict: dict, meet_nicknames: dict):
        lines = []
        for i, (name, nickname) in enumerate(zip(name_dict.values(), meet_nicknames.values())):
            if nickname is not None:
                lines.append(f"{i+1}. {name}: {nickname}")
        
        nicknames_list_text = "\n".join(lines)
        text = text.replace("[GOOGLE_MEET_NICKNAMES]", nicknames_list_text)
        return text

    def prepare_task_assigment_prompt(self, conversation_str: str, name_dict: dict, meet_nicknames: dict, json_log: dict = None):
        conversation_str = self.replace_speakers(conversation_str, name_dict)
        
        if json_log is not None:
            json_log["replaced_speakers_conversation"] = conversation_str

        task_assigment_prefix_modified = self.paste_date(copy.copy(self.task_assigment_prefix), json_log)
        task_assigment_prefix_modified = self.paste_google_meet_nicknames(task_assigment_prefix_modified, name_dict, meet_nicknames)

        content = task_assigment_prefix_modified + conversation_str
        
        return content
    
    def analyze_assignees_response(self, assignees_response: dict, json_log: dict):
        if json_log is not None:
            json_log["task_assigment_str"] = assignees_response
        
        if len(re.findall(r"\[CANT_HANDLE\]", assignees_response)) > 0:
            raise AICantHandleRequestException(assignees_response.replace("[CANT_HANDLE]", ""))
        
        try:
            assignee_dict = json.loads(assignees_response)
        except Exception as e:
            fix_content = self.fix_json_format_prefix + assignees_response
            fix_response = self.get_model_response(fix_content)
            assignee_dict = json.loads(fix_response)

        assignee_list: List[Assignee] = []
        for assignee_name, assagnee_data in assignee_dict.items():
            original_speaker_name = assagnee_data["original_speaker_name"]
            tasks = assagnee_data["tasks"]
            task_list: List[Task] = []
            for task_dict in tasks:
                approx_description = "-"
                try:
                    time = datetime.strptime(task_dict["time"], "%H:%M %d.%m.%Y")
                    local_tz = get_localzone()
                    time = time.replace(tzinfo=local_tz)
                except:
                    time = None
                    approx_description = task_dict["time"]
                
                title = task_dict["title"]
                deadline = Deadline(time, approx_description)
                task = Task(title, task_dict["task"], deadline)
                task_list.append(task)
            
            assignee = Assignee(assignee_name, original_speaker_name, task_list)
            assignee_list.append(assignee)
        
        if json_log is not None:
            json_log["assignees"] = [assignee.__dict__() for assignee in assignee_list]
        
        return assignee_list

    def add_assignee_to_bitrix(self, assignees: List[Assignee], speaker_to_user: dict):
        for assignee in assignees:
            user: BitrixUser = speaker_to_user[assignee.original_speaker_name]
            
            if user is not None:
                bitrix_tasks = [BitrixTask(title=task.title, created_by_id=1, responsible_id=user.id, discription=task.content, *({} if task.deadline.time is None else dict(deadline=task.deadline.time.isoformat()))) for task in assignee.tasks]
                for task in bitrix_tasks:
                    self.bitrix_manager.create_task_on_bitrix(task)
        

    def identify_assignee_for_participants(self, conversation_str: str, name_dict: dict, meet_nicknames: dict, json_log: dict = None):
        content = self.prepare_task_assigment_prompt(conversation_str, name_dict, meet_nicknames, json_log)
        response = self.get_model_response(content)

        assignee_list = self.analyze_assignees_response(response, json_log)
        return assignee_list

    def prepare_bitrix_finder_prompt(self, assignees: List[Assignee]):
        conversation_members_lines = []
        for i, assignee in enumerate(assignees):
            line = f"{i}: {assignee.original_speaker_name}"
            conversation_members_lines.append(line)
        conversation_members_str = "\n".join(conversation_members_lines)

        bitrix_finder_prompt = copy.copy(self.bitrix_finder_prompt)
        bitrix_finder_prompt = bitrix_finder_prompt.replace("[CONVERSATION_MEMBERS]", conversation_members_str)
        bitrix_finder_prompt = bitrix_finder_prompt.replace("[COMPANY MEMBERS]", self.employees_prompt)

        return bitrix_finder_prompt

    def analyze_bitrix_finder_prompt(self, response: str) -> dict:
        data: dict = json.loads(response)

        speaker_to_user = {}
        for speaker, value in data.items():
            bitrix_user = None
            if value["is_company_member"]:
                bitrix_user = self.full_names_to_bitrix_users[value["full_name"]]
            
            speaker_to_user[speaker] = bitrix_user
        
        return speaker_to_user

    def find_bitrix_full_name(self, assignees: List[Assignee]):
        speaker_to_user = {}
        for assignee in assignees:
            user = self.users_database.find_user(assignee.original_speaker_name, max_distance=0.7)
            if user is not None:
                assignee.original_speaker_name = f"{user.name} {user.second_name} {user.last_name}" + f" (изначально {assignee.name})"
                assignee.name = f"{user.name} {user.second_name} {user.last_name}"
            speaker_to_user[assignee.original_speaker_name] = user

        return speaker_to_user

    def modify_assignee_with_users(self, assignees: List[Assignee], speaker_to_user: dict):
        for assignee in assignees:
            if speaker_to_user[assignee.original_speaker_name] is not None:
                user = speaker_to_user[assignee.original_speaker_name]
                assignee.name = f"{user.name} {user.second_name} {user.last_name}"
        
        return assignees

    def generate_assignees(self, conversation_str: str, json_log: dict = None, preloaded_names: List[str] = [], identified_names_handler: IdentifiedNamesHandlerInterface = None) -> List[Assignee]: 
        name_dict, meet_nicknames = self.identify_assignee_names(conversation_str, json_log, preloaded_names=preloaded_names)
        assignee_list = self.identify_assignee_for_participants(conversation_str, name_dict, meet_nicknames, json_log)

        return assignee_list
    
    def correct_assignees_with_bitirx(self, assignee_list: List[Assignee]) -> dict:
        speaker_to_user = self.find_bitrix_full_name(assignee_list)
        self.modify_assignee_with_users(assignee_list, speaker_to_user)
        return speaker_to_user
    
    def generate_docx(self, conversation_str: str, template_path="docx_templates/default.docx", preloaded_names: List[str] = [], json_log: dict = None, identified_names_handler: IdentifiedNamesHandlerInterface = None, assignees: List[Assignee] = []):
        if len(assignees) == 0:
            assignees = self.generate_assignees(conversation_str, json_log, preloaded_names=preloaded_names, identified_names_handler=identified_names_handler)
        doc = self.get_docx_from_assignees(assignees, template_path)
        return doc
    
    def paste_preloaded_names(self, text: str, preloaded_names: List[str] = []):
        preloaded_names_list = "\n".join([f"{i+1}. {name}" for i, name in enumerate(preloaded_names)])
        preloaded_names_assigment = text.replace("[PRELOADED_NAMES]", preloaded_names_list)
        return preloaded_names_assigment

    def prepare_preloaded_names_assigment(self, preloaded_names: List[str] = []):
        preloaded_names_assigment = ""
        if len(preloaded_names) > 0:
            preloaded_names_assigment = self.paste_preloaded_names(copy.copy(self.preloaded_names_assigment), preloaded_names)
        return preloaded_names_assigment

    def prepare_name_identification_prompt(self, conversation_str: str, preloaded_names: List[str] = []):
        preloaded_names_assigment = self.prepare_preloaded_names_assigment(preloaded_names)
        
        content = self.name_identification_prefix + conversation_str
        content = content.replace("[PRELOADED_NAMES_PREFIX]", preloaded_names_assigment)
        return content

    def analyze_names_response(self, names_response: str, json_log: dict = None):
        raw_name_dict = yaml.safe_load(names_response)
        raw_name_dict = {} if raw_name_dict is None else raw_name_dict

        nicknames_pattern = r"\[(.*?)\]"
        name_pattern = r"^(.*?)(?:\[|$)"
        meet_nicknames = {}
        name_dict = {}

        for speaker, name in raw_name_dict.items():
            if name is not None:
                nickname = None
                matches = re.findall(nicknames_pattern, name)
                if len(matches) > 0:
                    nickname = matches[0].strip()
                
                meet_nicknames[speaker] = nickname
                name_dict[speaker] = re.findall(name_pattern, name)[0].strip()

        if json_log is not None:
            json_log["name_identification_str"] = names_response
            json_log["name_identification"] = raw_name_dict
        
        return name_dict, meet_nicknames

    def identify_assignee_names(self, conversation_str: str, json_log: dict = None, preloaded_names: List[str] = []) -> dict:
        content = self.prepare_name_identification_prompt(conversation_str, preloaded_names)
        response = self.get_model_response(content)

        name_dict, meet_nicknames = self.analyze_names_response(response, json_log)
        return name_dict, meet_nicknames
        
    
    def map_speaker_names(self, conversation_str: str):
        content = self.name_identification_prefix + conversation_str
        mapping_str = self.get_model_response(content)

        speakers = set(re.findall(r"\[-SPEAKER_\d+-\]", conversation_str))
        mapping = dict(re.findall(r"(\[-SPEAKER_\d+-\]):\s*(.*)", mapping_str))
        mapping = {speaker: name.strip() for speaker, name in mapping.items()}

        speakers_dict = {speaker: mapping.get(speaker, "Unknown") for speaker in speakers}

        return speakers_dict
    
    def add_task_table(self, doc, assignee):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'EloquityTableStyle'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Задача'
        hdr_cells[1].text = 'Крайний срок'

        for task in assignee.tasks:
            row_cells = table.add_row().cells

            row_cells[0].text = task.content
            row_cells[1].paragraphs[0].paragraph_format.alignment = 1

            time = task.deadline.time.strftime("%H:%M") + "\n" if task.deadline.time is not None else ""
            date = task.deadline.time.strftime("%d.%m.%Y") if task.deadline.time is not None else task.deadline.approx_discription 

            run0 = row_cells[1].paragraphs[0].add_run(time)
            run0.font.size = Pt(16)

            run1 = row_cells[1].paragraphs[0].add_run(date)
            run1.font.size = Pt(12)

        doc.add_paragraph('\n')
    
    def get_docx_from_assignees(self, assignees, template_path="docx_templates/default.docx"):
        doc = Document(template_path)

        doc.add_heading('Задачи', 1)

        for assignee in assignees:
            doc.add_heading(assignee.name, 2)
            self.add_task_table(doc, assignee)

        return doc