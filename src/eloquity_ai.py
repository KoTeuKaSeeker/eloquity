import requests
import re
from datetime import datetime, timedelta
from typing import List
import yaml
from docx import Document
from docx.shared import Pt
import json
import copy
from src.exeptions.ai_cant_handle_request_exception import AICantHandleRequestException


class Deadline():
    time: datetime
    approx_discription: str

    def __init__(self, time: datetime = None, approx_discription: str = "-"):
        self.time = time
        self.approx_discription = approx_discription
    
    def __str__(self):
        if self.time is None:
            return self.approx_discription
        else:
            return self.time.strftime("%Y-%m-%d %H:%M:%S")



class Task:
    def __init__(self, content: str, deadline: Deadline):
        self.content = content
        self.deadline = deadline
    
    def __str__(self):
        return f"({self.deadline}) {self.content}"
    
    def __dict__(self) -> dict:
        return {
            "content": self.content,
            "deadline": str(self.deadline)
        }


class Assignee:
    def __init__(self, name: str, tasks: List[Task]):
        self.name = name
        self.tasks = tasks
    
    def __str__(self):
        output = f"{self.name}:\n"
        for task in self.tasks:
            output += f"\t{task}\n"
        return output
    
    def __dict__(self) -> dict:
        return {
            "name": self.name,
            "tasks": [task.__dict__() for task in self.tasks]
        }



class EloquityAI:
    def __init__(self, api_key: str, model_name: str = 'gpt-4o'):
        self.api_url = "https://gptunnel.ru/v1/chat/completions"
        self.api_key = api_key
        self.name_identification_prefix = self._load_prefix("prefixes/name_identification.txt")
        self.task_assigment_prefix = self._load_prefix("prefixes/task_assigment.txt")
        self.fix_json_format_prefix = self._load_prefix("prefixes/fix_json_format.txt")
        self.model_name = model_name

    def _load_prefix(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except FileNotFoundError:
            print(f"Error: File '{file_path}' not found.")
            return ""

    def get_model_response(self, message):
        headers = {
            "Authorization": self.api_key,
            "Content-Type": "application/json"
        }
        data = {
            "model": self.model_name,
            "max_tokens": 2000,
            "messages": [{"role": "user", "content": message}]
        }

        try:
            response = requests.post(self.api_url, headers=headers, json=data)
            response.raise_for_status()  # Raise an error for HTTP errors
        except requests.exceptions.RequestException as e:
            print(f"Error: {e}")
        
        content = response.json()["choices"][0]["message"]["content"]

        return content
    
    def get_delta_time_from_str(self, time_string):
        pattern = r"(?:(\d+)d)?(?:(\d+)h)?(?:(\d+)m)?(?:(\d+)s)?"
        
        matches = re.match(pattern, time_string)
        
        if matches:
            days = int(matches.group(1) or 0)
            hours = int(matches.group(2) or 0)
            minutes = int(matches.group(3) or 0)
            seconds = int(matches.group(4) or 0)
            
            return timedelta(days=days, hours=hours, minutes=minutes, seconds=seconds)
        
        return timedelta()
    
    
    def generate_assignees(self, conversation_str: str, json_log: dict = None) -> List[Assignee]:
        if json_log is not None:
            json_log["original_conversation"] = conversation_str

        name_dict = self.identify_assignee_names(conversation_str, json_log)

        def replace_speakers(text, speakers):
            pattern = re.compile(r'\b(speaker_\d+)\b')  # Match words like speaker_0
            return pattern.sub(lambda match: f"[{speakers.get(match.group(1), match.group(1))}]", text)
        
        conversation_str = replace_speakers(conversation_str, name_dict)

        if json_log is not None:
            json_log["replaced_speakers_conversation"] = conversation_str
        
        current_date = datetime.now()
        current_time = "Текущая дата: " + current_date.strftime("%H:%M %d.%m.%Y") + "\n"
        task_assigment_prefix_with_current_time = copy.copy(self.task_assigment_prefix).replace("[CURRENT_DATE]", current_time)

        if json_log is not None:
            json_log["current_time_str"] = current_time

        content = task_assigment_prefix_with_current_time + conversation_str
        response = self.get_model_response(content)

        if json_log is not None:
            json_log["task_assigment_str"] = response
        
        if len(re.findall(r"\[CANT_HANDLE\]", response)) > 0:
            raise AICantHandleRequestException(response.replace("[CANT_HANDLE]", ""))
        
        # assignee_dict = yaml.safe_load(response)
        try:
            assignee_dict = json.loads(response)
        except Exception as e:
            fix_content = self.fix_json_format_prefix + response
            fix_response = self.get_model_response(fix_content)
            assignee_dict = json.loads(fix_response)

        current_datetime = datetime.now()

        assignee_list: List[Assignee] = []
        for assignee_name, tasks in assignee_dict.items():
            task_list: List[Task] = []
            for task_dict in tasks:
                # delta = self.get_delta_time_from_str(task_dict["time"])
                # task = Task(task_dict["task"], current_datetime + delta)
                approx_description = "-"
                try:
                    time = datetime.strptime(task_dict["time"], "%H:%M %d.%m.%Y")
                except:
                    time = None
                    approx_description = task_dict["time"]
                
                deadline = Deadline(time, approx_description)
                task = Task(task_dict["task"], deadline)
                task_list.append(task)
            
            assignee = Assignee(assignee_name, task_list)
            assignee_list.append(assignee)
        
        if json_log is not None:
            json_log["assignees"] = [assignee.__dict__() for assignee in assignee_list]

        return assignee_list
    
    def generate_docx(self, conversation_str: str, template_path="docx_templates/default.docx", json_log: dict = None):
        assignees = self.generate_assignees(conversation_str, json_log)

        # print("\n".join(str(assigne) for assigne in assignees))

        doc = self.get_docx_from_assignees(assignees, template_path)
        return doc
    
    def identify_assignee_names(self, conversation_str: str, json_log: dict = None) -> dict:
        content = self.name_identification_prefix + conversation_str
        response = self.get_model_response(content)
        name_dict = yaml.safe_load(response)

        if json_log is not None:
            json_log["name_identification_str"] = response
            json_log["name_identification"] = name_dict

        return name_dict
        
    
    def map_speaker_names(self, conversation_str: str):
        content = self.name_identification_prefix + conversation_str
        mapping_str = self.get_model_response(content)

        speakers = set(re.findall(r"\[-SPEAKER_\d+-\]", conversation_str))
        mapping = dict(re.findall(r"(\[-SPEAKER_\d+-\]):\s*(.*)", mapping_str))
        mapping = {speaker: name.strip() for speaker, name in mapping.items()}

        speakers_dict = {speaker: mapping.get(speaker, "Unknown") for speaker in speakers}

        return speakers_dict
    
    def add_task_table(self, doc, assignee):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'EloquityTableStyle'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Задача'
        hdr_cells[1].text = 'Крайний срок'

        for task in assignee.tasks:
            row_cells = table.add_row().cells

            row_cells[0].text = task.content
            row_cells[1].paragraphs[0].paragraph_format.alignment = 1

            time = task.deadline.time.strftime("%H:%M") + "\n" if task.deadline.time is not None else ""
            date = task.deadline.time.strftime("%d.%m.%Y") if task.deadline.time is not None else task.deadline.approx_discription 

            run0 = row_cells[1].paragraphs[0].add_run(time)
            run0.font.size = Pt(16)

            run1 = row_cells[1].paragraphs[0].add_run(date)
            run1.font.size = Pt(12)

        doc.add_paragraph('\n')
    
    def get_docx_from_assignees(self, assignees, template_path="docx_templates/default.docx"):
        doc = Document(template_path)

        doc.add_heading('Задачи', 1)

        for assignee in assignees:
            doc.add_heading(assignee.name, 2)
            self.add_task_table(doc, assignee)

        return doc