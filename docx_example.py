
from dotenv import load_dotenv
import os
from src.eloquity_ai import EloquityAI, Task
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH  # For alignment constants
from typing import List


def add_task_table(doc, assignee):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Задача'
    hdr_cells[1].text = 'Крайник срок'

    for task in assignee.tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = task.content
        row_cells[1].text = task.deadline
    
    header_cells = table.rows[0].cells
    for cell in header_cells:
        cell.paragraphs[0].style = 'Normal'  # Apply header style to the text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')


if __name__ == "__main__":
    load_dotenv()
    api_key = os.getenv("API_KEY")

    with open("conversation.txt", "r", encoding="utf-8") as file:
        conversation = file.read()

    model = EloquityAI(api_key=api_key)

    assignees = model.generate_assignees(conversation)

    doc = Document()
    heading_1_style = doc.styles['Heading 1']
    font = heading_1_style.font
    font.size = Pt(20)
    font.bold = False
    font.name = 'Montserrat'
    font.color.rgb = RGBColor(0, 0, 0)

    heading_1_style.paragraph_format.alignment = 1
    
    heading = doc.add_heading('Задачи', 1)

    heading_2_style = doc.styles['Heading 2']
    font = heading_2_style.font
    font.size = Pt(13)
    font.bold = False
    font.name = 'Montserrat'
    font.color.rgb = RGBColor(0, 0, 0)

    normal_style = doc.styles['Normal']
    font = normal_style.font
    font.size = Pt(10)
    font.bold = False
    font.name = 'Montserrat'
    font.color.rgb = RGBColor(0, 0, 0)

    for assignee in assignees:
        para = doc.add_heading(assignee.name, 2)
        add_task_table(doc, assignee)

    doc.save('example.docx')